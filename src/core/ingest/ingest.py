"""Document ingestion module for Sentio RAG system.

This module provides a robust, configurable interface for ingesting documents
into the vector database with proper error handling and logging.
"""

import logging
from pathlib import Path
from typing import Any

from src.core.chunking import ChunkingStrategy, TextChunker
from src.core.embeddings import BaseEmbedder, get_embedder
from src.core.models.document import Document
from src.core.vector_store import get_async_vector_store, get_vector_store

# Configure logging
logger = logging.getLogger(__name__)


class DocumentIngestor:
    """High-performance document ingestion with error handling and monitoring.

    This class handles the complete document ingestion pipeline:
    1. Loading documents from disk
    2. Chunking documents into smaller pieces
    3. Generating embeddings for chunks
    4. Storing chunks and embeddings in the vector database

    Args:
        collection_name: Name of the vector store collection
        chunk_size: Size of text chunks in tokens
        chunk_overlap: Overlap between chunks in tokens
        chunking_strategy: Strategy for chunking documents
        embedder_name: Name of the embedding provider to use
        vector_store_name: Name of the vector store backend
    """

    def __init__(
        self,
        collection_name: str = "Sentio_docs",
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        chunking_strategy: str = "recursive",
        embedder_name: str = "jina",
        vector_store_name: str = "qdrant",
    ) -> None:
        """Initialize the document ingestor with specified parameters."""
        self.collection_name = collection_name
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.chunking_strategy = chunking_strategy
        self.embedder_name = embedder_name
        self.vector_store_name = vector_store_name

        # Components will be initialized in initialize()
        self.chunker: TextChunker | None = None
        self.embedder: BaseEmbedder | None = None
        self.vector_store: Any | None = None
        self._initialized: bool = False

        # Stats tracking
        self._stats: dict[str, Any] = {
            "documents_processed": 0,
            "chunks_created": 0,
            "embeddings_generated": 0,
            "bytes_processed": 0,
        }

        logger.info("Document ingestor initialized with:")
        logger.info(f"  - Collection: {collection_name}")
        logger.info(f"  - Chunk size: {chunk_size}")
        logger.info(f"  - Chunk overlap: {chunk_overlap}")
        logger.info(f"  - Chunking strategy: {chunking_strategy}")
        logger.info(f"  - Embedder: {embedder_name}")
        logger.info(f"  - Vector store: {vector_store_name}")

    async def initialize(self) -> None:
        """Initialize all components with proper error handling.

        This method initializes the chunker, embedder, and vector store components.
        It should be called before using the ingestor.

        Raises:
            Exception: If any component fails to initialize.
        """
        try:
            logger.info("Initializing document ingestor components...")

            # Initialize chunker
            self.chunker = TextChunker(
                strategy=ChunkingStrategy(self.chunking_strategy),
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
            )
            logger.info("✓ Text chunker initialized")

            # Initialize embedder
            self.embedder = get_embedder(name=self.embedder_name)
            logger.info(f"✓ Embedding model initialized: {self.embedder_name}")

            # Initialize async vector store
            vector_size = self.embedder.dimension
            self.vector_store = await get_async_vector_store(
                name=self.vector_store_name,
                collection_name=self.collection_name,
                vector_size=vector_size,
            )
            logger.info(f"✓ Vector store initialized: {self.vector_store_name}")

            # Verify health
            if hasattr(self.vector_store, "health_check"):
                is_healthy = await self.vector_store.health_check()
                if is_healthy:
                    logger.info("✓ Vector store connection verified")
                else:
                    logger.warning("⚠ Vector store connection check failed")

            self._initialized = True
            logger.info("✅ Document ingestor fully initialized")

        except Exception as e:
            logger.error(f"Failed to initialize components: {e}")
            raise

    async def initialize_with_shared_components(self, embedder: Any, vector_store: Any) -> None:
        """Initialize with shared components to avoid creating new connections.

        This method uses pre-initialized shared components instead of creating new ones.
        This helps avoid creating multiple database connections and improves startup time.

        Args:
            embedder: Pre-initialized embedder instance
            vector_store: Pre-initialized async vector store instance

        Raises:
            Exception: If initialization fails.
        """
        try:
            logger.info("Initializing document ingestor with shared components...")

            # Initialize chunker
            self.chunker = TextChunker(
                strategy=ChunkingStrategy(self.chunking_strategy),
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
            )
            logger.info("✓ Text chunker initialized")

            # Use shared embedder
            self.embedder = embedder
            logger.info(f"✓ Embedding model initialized: {self.embedder_name}")

            # Use shared async vector store
            self.vector_store = vector_store
            logger.info(f"✓ Vector store initialized: {self.vector_store_name}")

            # Verify health
            if hasattr(self.vector_store, "health_check"):
                is_healthy = await self.vector_store.health_check()
                if is_healthy:
                    logger.info("✓ Vector store connection verified")
                else:
                    logger.warning("⚠ Vector store connection check failed")

            self._initialized = True
            logger.info("✅ Document ingestor fully initialized with shared components")

        except Exception as e:
            logger.error(f"Failed to initialize components with shared instances: {e}")
            raise

    def _read_file_content(self, file_path: Path) -> str:
        """Read content from a file with support for various formats.
        
        Args:
            file_path: Path to the file to read
            
        Returns:
            String content of the file
            
        Raises:
            ValueError: If the file format is unsupported or reading fails
        """
        suffix = file_path.suffix.lower()

        # Handle text files
        if suffix in {".txt", ".md", ".json", ".yaml", ".yml", ".html", ".css", ".js", ".py"}:
            return file_path.read_text(errors="replace")

        # Handle PDF files
        if suffix == ".pdf":
            try:
                # Import PyPDF2 here to avoid dependency if not needed
                from PyPDF2 import PdfReader

                reader = PdfReader(file_path)
                pages = [page.extract_text() or "" for page in reader.pages]
                return "\n\n".join(pages)
            except ImportError:
                logger.warning("PyPDF2 not installed. Cannot process PDF files.")
                raise ValueError("PyPDF2 not installed. Cannot process PDF files.")
            except Exception as e:
                logger.warning(f"Error reading PDF file {file_path}: {e}")
                raise ValueError(f"Error reading PDF file: {e}")

        # Handle DOCX files
        elif suffix == ".docx":
            try:
                # Import python-docx here to avoid dependency if not needed
                import docx

                doc = docx.Document(file_path)
                return "\n\n".join(paragraph.text for paragraph in doc.paragraphs)
            except ImportError:
                logger.warning("python-docx not installed. Cannot process DOCX files.")
                raise ValueError("python-docx not installed. Cannot process DOCX files.")
            except Exception as e:
                logger.warning(f"Error reading DOCX file {file_path}: {e}")
                raise ValueError(f"Error reading DOCX file: {e}")

        # Unsupported format
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    def _load_documents_from_directory(self, data_dir: Path) -> list[Document]:
        """Load documents from a directory.

        Args:
            data_dir: Directory containing documents to load

        Returns:
            List of Document objects

        Raises:
            ValueError: If no files are found or no valid documents are loaded
        """
        if not data_dir.exists():
            raise ValueError(f"Directory does not exist: {data_dir}")

        if not any(data_dir.iterdir()):
            logger.warning(f"No files found in directory: {data_dir}")
            return []  # Return empty list instead of raising exception

        logger.info(f"Loading documents from {data_dir}")

        documents: list[Document] = []
        supported_extensions = {".txt", ".md", ".pdf", ".docx", ".html", ".json", ".yaml", ".yml"}
        processed_files: set[str] = set()
        documents_count = 0

        # Process all files in the directory
        for file_path in data_dir.glob("**/*"):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    # Skip already processed files
                    if str(file_path) in processed_files:
                        continue

                    # Read file content using the appropriate method
                    content = self._read_file_content(file_path)

                    # Create document
                    doc = Document(
                        text=content,
                        metadata={
                            "source": str(file_path),
                            "filename": file_path.name,
                            "extension": file_path.suffix.lower(),
                            "size_bytes": file_path.stat().st_size,
                        }
                    )
                    documents.append(doc)
                    processed_files.add(str(file_path))

                    # Update stats
                    self._stats["bytes_processed"] += file_path.stat().st_size
                    documents_count += 1

                except Exception as e:
                    logger.warning(f"Error loading file {file_path}: {e}")

        if not documents:
            raise ValueError(f"No valid documents loaded from {data_dir}")

        # Update the counter directly, don't rely on previous value
        self._stats["documents_processed"] = documents_count
        logger.info(f"✓ Loaded {documents_count} documents")

        return documents

    async def _generate_embeddings(self, chunks: list[Document]) -> dict[str, list[float]]:
        """Generate embeddings for document chunks.

        Args:
            chunks: List of document chunks

        Returns:
            Dictionary mapping document IDs to embedding vectors

        Raises:
            Exception: If embedding generation fails
        """
        if not self.embedder:
            raise ValueError("Embedder not initialized. Call initialize() first.")

        logger.info(f"Generating embeddings for {len(chunks)} chunks")

        # Extract text from chunks
        texts = [chunk.text for chunk in chunks]

        # Generate embeddings
        try:
            # Use embed_async_many method which should be available on all BaseEmbedder implementations
            embeddings = await self.embedder.embed_async_many(texts)

            # Create mapping from document ID to embedding
            doc_embeddings = {
                chunks[i].id: embedding
                for i, embedding in enumerate(embeddings)
            }

            self._stats["embeddings_generated"] += len(embeddings)
            logger.info(f"✓ Generated {len(embeddings)} embeddings")

            return doc_embeddings

        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            raise

    async def _store_chunks_with_embeddings(
        self,
        chunks: list[Document],
        embeddings: dict[str, list[float]]
    ) -> None:
        """Store chunks and their embeddings in the vector database.

        Args:
            chunks: List of document chunks
            embeddings: Dictionary mapping document IDs to embedding vectors

        Raises:
            Exception: If storing in the vector database fails
        """
        if not self.vector_store:
            raise ValueError("Vector store not initialized. Call initialize() first.")

        logger.info(f"Storing {len(chunks)} chunks in vector database")

        try:
            # Prepare data for async vector store
            texts = []
            chunk_embeddings = []
            metadatas = []
            ids = []
            
            for chunk in chunks:
                # Skip if no embedding for this chunk
                if chunk.id not in embeddings:
                    logger.warning(f"No embedding found for chunk {chunk.id}, skipping")
                    continue

                texts.append(chunk.text)
                chunk_embeddings.append(embeddings[chunk.id])
                metadatas.append(chunk.metadata)
                ids.append(chunk.id)

            # Store using async vector store method
            if hasattr(self.vector_store, "add_embeddings"):
                logger.debug("Storing chunk IDs: %s", ids)
                await self.vector_store.add_embeddings(
                    texts=texts,
                    embeddings=chunk_embeddings,
                    metadatas=metadatas,
                    ids=ids,
                )
            else:
                # Fallback for other vector stores
                raise NotImplementedError(
                    f"Async storage method not implemented for {self.vector_store_name}"
                )

            logger.info(f"✓ Stored {len(texts)} chunks in vector database")

        except Exception as e:
            logger.error(f"Error storing chunks in vector database: {e}")
            raise

    async def ingest_documents(self, data_dir: str | Path) -> dict[str, Any]:
        """Ingest documents from a directory into the vector database.

        This method handles the complete ingestion pipeline:
        1. Loading documents from disk
        2. Chunking documents
        3. Generating embeddings
        4. Storing in vector database

        Args:
            data_dir: Directory containing documents to ingest

        Returns:
            Dictionary with ingestion statistics

        Raises:
            Exception: If ingestion fails at any step
        """
        # Convert to Path if string
        if isinstance(data_dir, str):
            data_dir = Path(data_dir)

        logger.info(f"🔍 Starting document ingestion from: {data_dir}")

        # Ensure components are initialized
        if not all([self.chunker, self.embedder, self.vector_store]):
            await self.initialize()

        try:
            # Load documents
            documents = self._load_documents_from_directory(data_dir)

            # If documents were loaded but the count is still zero, update it
            # This is needed for tests that mock _load_documents_from_directory
            if len(documents) > 0 and self._stats["documents_processed"] == 0:
                self._stats["documents_processed"] = len(documents)

            # Chunk documents
            assert self.chunker is not None  # For type checking
            chunks = self.chunker.split(documents)
            self._stats["chunks_created"] += len(chunks)
            logger.info(f"✓ Created {len(chunks)} chunks from {len(documents)} documents")

            # Generate embeddings
            embeddings = await self._generate_embeddings(chunks)

            # Store chunks with embeddings
            await self._store_chunks_with_embeddings(chunks, embeddings)

            logger.info("✅ Document ingestion completed successfully")
            logger.info(f"   Documents processed: {self._stats['documents_processed']}")
            logger.info(f"   Chunks created: {self._stats['chunks_created']}")
            logger.info(f"   Embeddings generated: {self._stats['embeddings_generated']}")
            logger.info(f"   Bytes processed: {self._stats['bytes_processed']}")

            return self._stats.copy()

        except Exception as e:
            logger.error(f"Document ingestion failed: {e}")
            raise Exception(f"Document ingestion failed: {e}") from e

    @property
    def stats(self) -> dict[str, Any]:
        """Get current ingestion statistics."""
        return self._stats.copy()

    async def ingest_document(self, document: Document) -> dict[str, Any]:
        """Ingest a single in-memory document end-to-end.

        Args:
            document: Document to chunk, embed and store.

        Returns:
            Minimal stats for the single-document ingestion.
        """
        # Ensure components are initialized
        if not all([self.chunker, self.embedder, self.vector_store]):
            await self.initialize()

        # Chunk
        assert self.chunker is not None
        chunks = self.chunker.split([document])

        # Embed
        embeddings = await self._generate_embeddings(chunks)

        # Store
        await self._store_chunks_with_embeddings(chunks, embeddings)

        return {
            "documents_processed": 1,
            "chunks_created": len(chunks),
            "embeddings_generated": len(embeddings),
            "bytes_processed": len(document.text.encode("utf-8", errors="ignore")),
        }


async def ingest_directory(
    data_dir: str | Path,
    collection_name: str = "Sentio_docs",
    chunk_size: int = 512,
    chunk_overlap: int = 64,
    chunking_strategy: str = "recursive",
    embedder_name: str = "jina",
    vector_store_name: str = "qdrant",
) -> dict[str, Any]:
    """Convenience function to ingest documents from a directory.

    This function creates a DocumentIngestor and runs the ingestion pipeline.

    Args:
        data_dir: Directory containing documents to ingest
        collection_name: Name of the vector store collection
        chunk_size: Size of text chunks in tokens
        chunk_overlap: Overlap between chunks in tokens
        chunking_strategy: Strategy for chunking documents
        embedder_name: Name of the embedding provider to use
        vector_store_name: Name of the vector store backend

    Returns:
        Dictionary with ingestion statistics

    Raises:
        Exception: If ingestion fails at any step
    """
    ingestor = DocumentIngestor(
        collection_name=collection_name,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        chunking_strategy=chunking_strategy,
        embedder_name=embedder_name,
        vector_store_name=vector_store_name,
    )

    await ingestor.initialize()
    return await ingestor.ingest_documents(data_dir)
